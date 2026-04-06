"""Generate .docx resumes for all 6 pending applications, matching existing resume formatting."""

from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Formatting constants (matched from Close resume)
FONT = 'Arial'
NAME_SIZE = Pt(18)
CONTACT_SIZE = Pt(8)
LINK_SIZE = Pt(7.5)
SECTION_HEADER_SIZE = Pt(9.5)
BODY_SIZE = Pt(8.5)
TOP_MARGIN = 365760
BOTTOM_MARGIN = 365760
LEFT_MARGIN = 502920
RIGHT_MARGIN = 502920


def create_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    return doc


def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(name)
    r.bold = True
    r.font.size = NAME_SIZE
    r.font.name = FONT


def add_contact_line(doc, text, is_link=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = LINK_SIZE if is_link else CONTACT_SIZE
    r.font.name = FONT


def add_spacer(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def add_section_header(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = True
    r.font.size = SECTION_HEADER_SIZE
    r.font.name = FONT


def add_body_text(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = BODY_SIZE
    r.font.name = FONT


def add_skill_line(doc, category, content):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r1 = p.add_run(f'{category}: ')
    r1.bold = True
    r1.font.size = BODY_SIZE
    r1.font.name = FONT
    r2 = p.add_run(content)
    r2.font.size = BODY_SIZE
    r2.font.name = FONT


def add_job_header(doc, title_company, dates_location):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(title_company)
    r.bold = True
    r.font.size = BODY_SIZE
    r.font.name = FONT
    p2 = doc.add_paragraph()
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(1)
    pf2.line_spacing = 1.0
    r2 = p2.add_run(dates_location)
    r2.font.size = BODY_SIZE
    r2.font.name = FONT


def add_project_header(doc, title, link=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(title)
    r.bold = True
    r.font.size = BODY_SIZE
    r.font.name = FONT
    if link:
        p2 = doc.add_paragraph()
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(0)
        pf2.line_spacing = 1.0
        r2 = p2.add_run(link)
        r2.font.size = LINK_SIZE
        r2.font.name = FONT


def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.25)
    r = p.add_run(f'\u2022 {text}')
    r.font.size = BODY_SIZE
    r.font.name = FONT


def add_education(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        r = p.add_run(line)
        r.font.size = BODY_SIZE
        r.font.name = FONT


def add_header_block(doc, location='Cincinnati, OH | Open to Remote'):
    add_name(doc, 'ISAIAH WASHINGTON')
    add_contact_line(doc, location)
    add_contact_line(doc, '513.544.9022 | isaiahwashington48@gmail.com')
    add_contact_line(doc, 'linkedin.com/in/isaiah-washington48 | github.com/Izayauh/whisper', is_link=True)


# ─── NEXUS WIFI BULLETS (variations) ───

NEXUS_6 = [
    'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and device failures',
    'Deliver 1,500+ hours of 100% remote customer-facing support, guiding users through troubleshooting and resolution',
    'Diagnose and triage issues with access points, switches, gateways, and modems, coordinating hardware replacement when needed',
    'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope',
    'Manage active tickets using an in-house platform, including outbound follow-up calls and status updates',
    'Resolve an estimated 5\u201310 tickets per morning shift and 2\u20135 per night shift through remote troubleshooting and coordination',
]

NEXUS_5_IMPL = [
    'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and device failures',
    'Deliver 1,500+ hours of 100% remote customer-facing support, guiding users through troubleshooting and resolution',
    'Diagnose and triage issues with access points, switches, gateways, and modems, coordinating hardware replacement when needed',
    'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope',
    'Perform camera setup and implementation work alongside core network support responsibilities',
]

NEXUS_4_SE = [
    'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and device failures',
    'Deliver 1,500+ hours of 100% remote customer-facing support, guiding users through troubleshooting and resolution',
    'Diagnose and triage issues with access points, switches, gateways, and modems, coordinating hardware replacement when needed',
    'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope',
]

# ─── IMPULSE BULLETS (variations) ───

IMPULSE_2 = [
    'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub',
    'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution',
]

IMPULSE_3_EVE = [
    'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub (57 commits, 2 releases)',
    'Implemented NVIDIA CUDA acceleration with smart model selection optimizing speed/quality based on dictation length',
    'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution',
]

IMPULSE_4 = [
    'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub (57 commits, 2 releases)',
    'Implemented NVIDIA CUDA acceleration with smart model selection optimizing speed/quality based on dictation length',
    'Built a product landing page using React, TypeScript, Tailwind CSS, and Vercel with working beta signup flow and interactive demo',
    'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution',
]

JARVIS_2 = [
    'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 tool functions',
    'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1',
]

JARVIS_4 = [
    'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 tool functions',
    'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1',
    'Implemented agent specialization across planning, execution, verification, and reporting phases with structured handoffs between agents',
    'Built crash recovery, process lifecycle management, and idempotent operations for reliable automated execution',
]


# ═══════════════════════════════════════════════════════════════
# 1. SAMSARA — Technical Support Engineer
# ═══════════════════════════════════════════════════════════════
def build_samsara():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote experience triaging hardware and software issues across 150+ managed properties. 1,500+ hours of direct customer interaction resolving connectivity outages, equipment failures, and infrastructure problems through structured diagnostic workflows. Comfortable translating complex technical concepts into clear guidance for non-technical users. Google IT Support certified. Builds Python-based tools independently, including a shipped dictation app with GPU acceleration. Seeking a Technical Support Engineer role supporting connected hardware and software at scale.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'Support & Troubleshooting', 'Hardware and software triage (access points, switches, gateways, modems, cameras), remote diagnostics, connectivity issue resolution, equipment failure coordination, structured escalation workflows')
    add_skill_line(doc, 'Operations', 'Ticket management, knowledge documentation, outbound follow-up, cross-functional coordination with engineering teams')
    add_skill_line(doc, 'Programming & Scripting', 'Python, JSON, SQL (coursework), HTML/CSS/JavaScript')
    add_skill_line(doc, 'AI & Tools', 'Google Gemini API, OpenAI Whisper, NVIDIA CUDA, Git, GitHub, Windows/WSL')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    for b in NEXUS_6:
        add_bullet(doc, b)
    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')
    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Assisted walk-in clients, scheduled appointments, and answered phone lines in a client-facing administrative role')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    for b in IMPULSE_2:
        add_bullet(doc, b)
    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    for b in JARVIS_2:
        add_bullet(doc, b)

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 2. EVE — Technical Support Specialist
# ═══════════════════════════════════════════════════════════════
def build_eve():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote experience resolving customer-facing issues across 150+ managed properties. 1,500+ hours of direct support through structured diagnostic workflows, with consistent ticket throughput under time-sensitive conditions. Builds AI-powered tools independently \u2014 including a shipped dictation app using OpenAI Whisper and a multi-agent voice assistant using Google Gemini API. Seeking a Technical Support Specialist role on an AI-native platform where fast ticket resolution and technical fluency drive customer outcomes.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'Support & Troubleshooting', 'Remote diagnostics, SaaS and infrastructure triage, connectivity issue resolution, equipment failure coordination, structured escalation workflows, ticket prioritization')
    add_skill_line(doc, 'Operations', 'Ticket management, feature request documentation, cross-functional coordination, outbound follow-up, system setup support')
    add_skill_line(doc, 'AI & Integration', 'Google Gemini API (function calling, real-time audio), OpenAI Whisper (local inference), NVIDIA CUDA, API integration, multi-agent architecture')
    add_skill_line(doc, 'Languages & Tools', 'Python, React, TypeScript, SQL (coursework), Git, GitHub, Vercel, Windows/WSL')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    for b in NEXUS_6:
        add_bullet(doc, b)
    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')
    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Assisted walk-in clients, scheduled appointments, and answered phone lines in a client-facing administrative role')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    for b in IMPULSE_3_EVE:
        add_bullet(doc, b)
    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    for b in JARVIS_2:
        add_bullet(doc, b)

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 3. AGORA — Customer Support Engineer
# ═══════════════════════════════════════════════════════════════
def build_agora():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote troubleshooting across 150+ client properties, combined with hands-on experience building AI-powered systems using LLMs and real-time audio. Designed and built a 6-agent voice-controlled assistant using Google Gemini API with structured agent handoffs, and shipped Impulse \u2014 a GPU-accelerated speech-to-text app using OpenAI Whisper. Seeking a Customer Support Engineer role supporting developers building with real-time communication and AI technologies.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'AI & LLM', 'Google Gemini API (function calling, real-time audio), OpenAI Whisper (speech-to-text, local inference), NVIDIA CUDA, multi-agent architecture, deterministic pipeline design, prompt engineering')
    add_skill_line(doc, 'Languages & Frameworks', 'Python, React, TypeScript, HTML/CSS/JavaScript, SQL (coursework), Tailwind CSS')
    add_skill_line(doc, 'Architecture', 'Multi-agent systems (PLAN/EXECUTE/VERIFY/REPORT), OSC protocol, crash recovery, idempotent operations, API integration')
    add_skill_line(doc, 'Support & Operations', 'Remote troubleshooting, ticket management, customer-facing communication, device triage, structured diagnostic workflows')
    add_skill_line(doc, 'Tools & Platforms', 'Git, GitHub, Vite, Vercel, Windows/WSL, virtual environments')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    for b in JARVIS_4:
        add_bullet(doc, b)
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    for b in IMPULSE_4:
        add_bullet(doc, b)

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    for b in NEXUS_4_SE:
        add_bullet(doc, b)
    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 4. ALPINE IQ — Implementation Specialist
# ═══════════════════════════════════════════════════════════════
def build_alpine_iq():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical professional who builds working systems \u2014 from a spreadsheet-based inventory tracker that improved dealership operations, to a multi-agent AI assistant with 62 tool functions and a deterministic execution pipeline. 4+ years of remote support experience across 150+ properties, with a focus on turning requirements into reliable technical solutions. Comfortable managing multiple accounts simultaneously and translating technical concepts into clear guidance for non-technical users. Seeking an Implementation Specialist role driving customer onboarding and platform adoption.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'System Building', 'Pipeline architecture (PLAN/EXECUTE/VERIFY/REPORT), multi-agent design, API integration, data mapping, spreadsheet-based tracking systems, process automation')
    add_skill_line(doc, 'Customer-Facing', 'Remote troubleshooting (1,500+ hours), customer onboarding guidance, multi-account management, cross-functional coordination, ticket workflow management')
    add_skill_line(doc, 'Languages & Frameworks', 'Python, React, TypeScript, HTML/CSS/JavaScript, SQL (coursework), Tailwind CSS')
    add_skill_line(doc, 'AI & Tools', 'Google Gemini API, OpenAI Whisper, NVIDIA CUDA, Git, GitHub, Vercel, Windows/WSL')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    for b in NEXUS_5_IMPL:
        add_bullet(doc, b)
    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Improved visibility and organization across dealership operations during a period of significant business growth')
    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    for b in IMPULSE_4:
        add_bullet(doc, b)
    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    for b in JARVIS_2:
        add_bullet(doc, b)

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        'Relevant coursework: Database Design & Development, Business Management, Small Business Operations',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 5. WELLTHY — Implementation Specialist
# ═══════════════════════════════════════════════════════════════
def build_wellthy():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical professional who builds working systems \u2014 from a spreadsheet-based inventory tracker that improved dealership operations, to a multi-agent AI assistant with 62 tool functions and a deterministic execution pipeline. 4+ years of remote support experience across 150+ properties, with a focus on coordinating cross-functional teams and turning requirements into reliable solutions. Comfortable with AI-powered tools and structured project workflows. Seeking an Implementation Specialist role driving client onboarding and operational setup.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'System Building', 'Pipeline architecture (PLAN/EXECUTE/VERIFY/REPORT), multi-agent design, API integration, spreadsheet-based tracking systems, process automation, project coordination')
    add_skill_line(doc, 'Customer-Facing', 'Remote troubleshooting (1,500+ hours), client onboarding support, cross-functional coordination, ticket workflow management, training and guidance')
    add_skill_line(doc, 'Languages & Frameworks', 'Python, React, TypeScript, HTML/CSS/JavaScript, SQL (coursework), Tailwind CSS')
    add_skill_line(doc, 'AI & Tools', 'Google Gemini API, OpenAI Whisper, NVIDIA CUDA, Git, GitHub, Vercel, Windows/WSL')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    for b in NEXUS_5_IMPL:
        add_bullet(doc, b)
    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Improved visibility and organization across dealership operations during a period of significant business growth')
    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    for b in IMPULSE_4:
        add_bullet(doc, b)
    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    for b in JARVIS_2:
        add_bullet(doc, b)

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        'Relevant coursework: Database Design & Development, Business Management, Small Business Operations',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 6. NEXTPATIENT — Customer Success Associate
# ═══════════════════════════════════════════════════════════════
def build_nextpatient():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical professional who builds working systems \u2014 from a spreadsheet-based inventory tracker that improved dealership operations, to a multi-agent AI assistant with 62 tool functions and a deterministic execution pipeline. 4+ years of remote support experience across 150+ properties, managing onboarding, configuration, and troubleshooting for diverse client sites. Seeking a Customer Success Associate role where hands-on product configuration and customer-facing communication drive client outcomes.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'System Building', 'Pipeline architecture (PLAN/EXECUTE/VERIFY/REPORT), multi-agent design, API integration, product configuration, spreadsheet-based tracking systems, process automation')
    add_skill_line(doc, 'Customer-Facing', 'Remote troubleshooting (1,500+ hours), client onboarding support, cross-functional coordination, ticket workflow management, customer training and feedback communication')
    add_skill_line(doc, 'Languages & Frameworks', 'Python, React, TypeScript, HTML/CSS/JavaScript, SQL (coursework), Tailwind CSS')
    add_skill_line(doc, 'AI & Tools', 'Google Gemini API, OpenAI Whisper, NVIDIA CUDA, Git, GitHub, Vercel, Windows/WSL')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    for b in NEXUS_5_IMPL:
        add_bullet(doc, b)
    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Assisted walk-in clients, scheduled appointments, and answered phone lines in a client-facing administrative role')
    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    for b in IMPULSE_4:
        add_bullet(doc, b)
    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    for b in JARVIS_2:
        add_bullet(doc, b)

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        'Relevant coursework: Database Design & Development, Business Management, Small Business Operations',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# GENERATE ALL
# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    base = r'C:\Users\isaia\Documents\Job_Project\04_Applications'

    builds = [
        ('Samsara', 'Resume_Samsara_TSE.docx', build_samsara),
        ('Eve', 'Resume_Eve_TSS.docx', build_eve),
        ('Agora', 'Resume_Agora_CSE.docx', build_agora),
        ('Alpine_IQ', 'Resume_Alpine_IQ_Implementation.docx', build_alpine_iq),
        ('Wellthy', 'Resume_Wellthy_Implementation.docx', build_wellthy),
        ('NextPatient', 'Resume_NextPatient_CSA.docx', build_nextpatient),
    ]

    for folder, filename, builder in builds:
        path = os.path.join(base, folder, filename)
        doc = builder()
        doc.save(path)
        print(f'Created: {path}')

    print('\nAll 6 resumes generated.')
