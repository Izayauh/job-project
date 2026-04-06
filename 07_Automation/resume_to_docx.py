"""
Resume Markdown-to-DOCX converter.
Reads the plain-text resume block from a tailored .md file and outputs a
clean, formatted .docx matching Isaiah's resume template style.

Usage:
    python resume_to_docx.py <input.md> [output.docx]
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def extract_resume_text(md_path: str) -> str:
    content = Path(md_path).read_text(encoding="utf-8")
    match = re.search(r'```\n(.*?)```', content, re.DOTALL)
    if not match:
        print("Error: Could not find a ``` code block in the file.")
        sys.exit(1)
    return match.group(1)


def is_special_line(line: str) -> bool:
    """Check if a line is a header, bullet, separator, URL, or other structural element."""
    s = line.strip()
    if not s:
        return True
    if re.match(r'^[A-Z][A-Z &/]+$', s):  # ALL CAPS header
        return True
    if s.startswith(('•', '*', '\u2022')):  # bullet
        return True
    if all(c in '-\u2500\u2501\u2502\u2503' for c in s):  # separator
        return True
    if s.startswith(('github.com', 'http', 'impulse-', 'linkedin.com')):  # URL
        return True
    if ':' in s and not s.endswith(':') and len(s) > 20:  # skills line
        colon_idx = s.index(':')
        if colon_idx < 30:
            return True
    if line.startswith('    '):  # indented continuation
        return True
    # Contact lines (contain | for separating info, or @ for email)
    if '|' in s:
        return True
    if '@' in s:
        return True
    # Lines with em-dashes — only if short AND starts with uppercase (job titles, not body text)
    if (len(s) < 80 and s[0].isupper() and
            ('\u2014' in s or '\u2013' in s or '—' in s or '–' in s)):
        return True
    # Date-prefixed lines (location/date under job titles)
    if re.match(r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})', s):
        return True
    # Short lines that are likely sub-headers (project names, school names, cert names)
    if len(s.split()) <= 10 and not s.endswith('.'):
        return True
    return False


def join_paragraph_lines(text: str) -> str:
    """Join consecutive plain body-text lines into single paragraphs."""
    lines = text.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # If this is a special/structural line, keep it as-is
        if is_special_line(line):
            result.append(line)
            i += 1
            continue
        # It's a body text line — collect consecutive body lines
        paragraph_parts = [line.strip()]
        while (i + 1 < len(lines) and
               not is_special_line(lines[i + 1]) and
               lines[i + 1].strip()):
            i += 1
            paragraph_parts.append(lines[i].strip())
        result.append(' '.join(paragraph_parts))
        i += 1
    return '\n'.join(result)


def set_run_font(run, name="Arial", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Force Arial for East Asian fallback too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def add_paragraph(doc, text, size=10, bold=False, italic=False,
                  align=None, color=None, space_after=Pt(1),
                  space_before=Pt(0), line_spacing=1.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_after = space_after
    fmt.space_before = space_before
    fmt.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_separator(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(2)
    # Add bottom border to paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'AAAAAA',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


BODY_COLOR = RGBColor(0x00, 0x1D, 0x35)  # dark navy from template

# Font sizes - tuned for dense content on one page
NAME_SIZE = 18
CONTACT_SIZE = 8
HEADER_SIZE = 9.5
BODY_SIZE = 8.5
BULLET_SIZE = 8.5
SUB_HEADER_SIZE = 8.5


def build_docx(text: str, output_path: str):
    # Pre-process: join continuation lines into paragraphs
    text = join_paragraph_lines(text)

    doc = Document()

    # Page setup - letter, tight margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines at the start
        if not stripped and i == 0:
            i += 1
            continue

        # Separator lines (box-drawing characters)
        if stripped and all(c in '-\u2500\u2501\u2502\u2503' for c in stripped):
            add_separator(doc)
            i += 1
            continue

        # Name line (first substantial line, all caps)
        if i <= 2 and stripped and re.match(r'^[A-Z][A-Z ]+$', stripped):
            add_paragraph(doc, stripped, size=NAME_SIZE, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(1), space_before=Pt(0))
            i += 1
            continue

        # Contact lines (near top, contain | or @)
        if i <= 5 and ('|' in stripped or '@' in stripped or
                       stripped.startswith('linkedin') or stripped.startswith('github')):
            add_paragraph(doc, stripped, size=CONTACT_SIZE,
                         align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(0), color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        # Section headers (ALL CAPS, like PROFESSIONAL SUMMARY, TECHNICAL SKILLS, etc.)
        if re.match(r'^[A-Z][A-Z &/]+$', stripped):
            add_paragraph(doc, stripped, size=HEADER_SIZE, bold=True,
                         space_before=Pt(4), space_after=Pt(2))
            i += 1
            continue

        # Sub-headers: job titles/project names (contain em-dash or date ranges)
        if ('\u2014' in stripped or '—' in stripped or
            ('–' in stripped and any(str(yr) in stripped for yr in range(2019, 2027))) or
            ('-' in stripped and 'Present' in stripped and any(str(yr) in stripped for yr in range(2019, 2027)))):
            add_paragraph(doc, stripped, size=SUB_HEADER_SIZE, bold=True,
                         space_before=Pt(3), space_after=Pt(0))
            i += 1
            continue

        # Project names (line before URL or bullets, italic)
        if (stripped and not stripped.startswith('•') and not stripped.startswith('*') and
            i + 1 < len(lines) and
            (lines[i+1].strip().startswith('github.com') or
             lines[i+1].strip().startswith('http') or
             lines[i+1].strip().startswith('impulse-') or
             lines[i+1].strip().startswith('•') or
             lines[i+1].strip().startswith('*'))):
            # Check if it looks like a project name (not a regular paragraph)
            if len(stripped.split()) <= 15 and not stripped.endswith('.'):
                add_paragraph(doc, stripped, size=SUB_HEADER_SIZE, bold=True,
                             space_before=Pt(3), space_after=Pt(0))
                i += 1
                continue

        # URL lines
        if stripped.startswith(('github.com', 'http', 'impulse-', 'linkedin.com')):
            add_paragraph(doc, stripped, size=BODY_SIZE - 1, italic=True,
                         color=RGBColor(0x55, 0x55, 0x55),
                         space_after=Pt(1))
            i += 1
            continue

        # Location/date lines (right after a job title)
        if stripped and re.match(r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})', stripped):
            add_paragraph(doc, stripped, size=BODY_SIZE,
                         color=RGBColor(0x55, 0x55, 0x55),
                         space_after=Pt(1))
            i += 1
            continue

        # Bullet points
        if stripped.startswith('•') or stripped.startswith('*'):
            # Collect continuation lines
            bullet_text = stripped
            while (i + 1 < len(lines) and
                   lines[i+1].startswith('    ') and
                   not lines[i+1].strip().startswith('•') and
                   not lines[i+1].strip().startswith('*') and
                   lines[i+1].strip()):
                i += 1
                bullet_text += ' ' + lines[i].strip()

            # Replace * with bullet character
            if bullet_text.startswith('* '):
                bullet_text = '\u2022 ' + bullet_text[2:]
            elif bullet_text.startswith('•'):
                pass  # already has bullet

            p = add_paragraph(doc, "", size=BULLET_SIZE,
                            space_after=Pt(0.5), space_before=Pt(0),
                            line_spacing=1.05)
            # Add indent
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            run = p.add_run(bullet_text)
            set_run_font(run, size=BULLET_SIZE, color=BODY_COLOR)
            i += 1
            continue

        # Skills lines (contain colon alignment like "Languages & Frameworks: ...")
        if ':' in stripped and not stripped.endswith(':') and len(stripped) > 20:
            # Split into label and value
            colon_idx = stripped.index(':')
            label = stripped[:colon_idx + 1]
            value = stripped[colon_idx + 1:].strip()

            # Collect continuation lines (indented)
            while (i + 1 < len(lines) and
                   lines[i+1].startswith('    ') and
                   not ':' in lines[i+1][:30] and
                   lines[i+1].strip()):
                i += 1
                value += ' ' + lines[i].strip()

            p = add_paragraph(doc, "", size=BODY_SIZE,
                            space_after=Pt(0.5), space_before=Pt(0),
                            line_spacing=1.05)
            p.paragraph_format.left_indent = Inches(0.15)
            run_label = p.add_run(label + ' ')
            set_run_font(run_label, size=BODY_SIZE, bold=True, color=BODY_COLOR)
            run_value = p.add_run(value)
            set_run_font(run_value, size=BODY_SIZE, color=BODY_COLOR)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Default text
        add_paragraph(doc, stripped, size=BODY_SIZE,
                     color=BODY_COLOR, space_after=Pt(1),
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        i += 1

    doc.save(output_path)
    print(f"DOCX saved: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python resume_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        output_path = str(Path(input_path).with_suffix('.docx'))

    text = extract_resume_text(input_path)
    build_docx(text, output_path)


if __name__ == "__main__":
    main()
