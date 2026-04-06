"""Generate .docx resumes for 4 Austin job applications, matching existing resume formatting."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Formatting constants (matched from existing generate_docx.py)
FONT = 'Arial'
NAME_SIZE = Pt(18)
CONTACT_SIZE = Pt(8)
LINK_SIZE = Pt(7.5)
SECTION_HEADER_SIZE = Pt(9.5)
BODY_SIZE = Pt(8.5)
TOP_MARGIN = 365760
BOTTOM_MARGIN = 365760
LEFT_MARGIN = 502920
RIGHT_MARGIN = 502920


def create_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    return doc


def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(name)
    r.bold = True
    r.font.size = NAME_SIZE
    r.font.name = FONT


def add_contact_line(doc, text, is_link=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = LINK_SIZE if is_link else CONTACT_SIZE
    r.font.name = FONT


def add_spacer(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def add_section_header(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = True
    r.font.size = SECTION_HEADER_SIZE
    r.font.name = FONT


def add_body_text(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = BODY_SIZE
    r.font.name = FONT


def add_skill_line(doc, category, content):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r1 = p.add_run(f'{category}: ')
    r1.bold = True
    r1.font.size = BODY_SIZE
    r1.font.name = FONT
    r2 = p.add_run(content)
    r2.font.size = BODY_SIZE
    r2.font.name = FONT


def add_job_header(doc, title_company, dates_location):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(title_company)
    r.bold = True
    r.font.size = BODY_SIZE
    r.font.name = FONT
    p2 = doc.add_paragraph()
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(1)
    pf2.line_spacing = 1.0
    r2 = p2.add_run(dates_location)
    r2.font.size = BODY_SIZE
    r2.font.name = FONT


def add_project_header(doc, title, link=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(title)
    r.bold = True
    r.font.size = BODY_SIZE
    r.font.name = FONT
    if link:
        p2 = doc.add_paragraph()
        pf2 = p2.paragraph_format
        pf2.space_before = Pt(0)
        pf2.space_after = Pt(0)
        pf2.line_spacing = 1.0
        r2 = p2.add_run(link)
        r2.font.size = LINK_SIZE
        r2.font.name = FONT


def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.25)
    r = p.add_run(f'\u2022 {text}')
    r.font.size = BODY_SIZE
    r.font.name = FONT


def add_education(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        r = p.add_run(line)
        r.font.size = BODY_SIZE
        r.font.name = FONT


def add_header_block(doc):
    add_name(doc, 'ISAIAH WASHINGTON')
    add_contact_line(doc, 'Cincinnati, OH | Open to Austin, TX and Remote')
    add_contact_line(doc, '513.544.9022 | isaiahwashington48@gmail.com')
    add_contact_line(doc, 'linkedin.com/in/isaiah-washington48 | github.com/Izayauh/whisper', is_link=True)


# ═══════════════════════════════════════════════════════════════
# 1. POSTMAN — Technical Support Engineer
# ═══════════════════════════════════════════════════════════════
def build_postman():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote experience triaging device and infrastructure issues across 150+ managed properties. 1,500+ hours of direct customer interaction resolving connectivity outages, equipment failures, and support issues through structured diagnostic workflows. Builds API-driven tools independently \u2014 shipped a GPU-accelerated dictation app and built a voice-controlled AI assistant integrating Google Gemini API with real-time function calling. Google IT Support certified. Seeking a Technical Support Engineer role at Postman where hands-on API experience and clear customer communication drive issue resolution.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'Support & Troubleshooting', 'Device and infrastructure triage (access points, switches, gateways, modems), remote diagnostics, connectivity issue resolution, equipment failure coordination, structured escalation workflows')
    add_skill_line(doc, 'APIs & Integration', 'Google Gemini API (function calling, real-time audio streaming), OSC protocol, REST-style API integration, multi-agent tool dispatch')
    add_skill_line(doc, 'Tools & Platforms', 'Git, GitHub, Jira-style ticketing workflows, in-house support platforms, Windows OS, Microsoft Teams, Vercel')
    add_skill_line(doc, 'Languages', 'Python, SQL (coursework), HTML/CSS/JavaScript, React, TypeScript')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    add_bullet(doc, 'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and device failures')
    add_bullet(doc, 'Deliver 1,500+ hours of 100% remote customer-facing support, guiding users through troubleshooting and resolution via phone, screen-sharing, and written follow-ups')
    add_bullet(doc, 'Diagnose and triage issues with access points, switches, gateways, and modems, coordinating hardware replacement when needed')
    add_bullet(doc, 'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope')
    add_bullet(doc, 'Manage active tickets using an in-house platform, including outbound follow-up calls and status updates')
    add_bullet(doc, 'Resolve an estimated 5\u201310 tickets per morning shift and 2\u20135 per night shift through remote troubleshooting and coordination')

    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Assisted walk-in clients, scheduled appointments, and answered phone lines in a client-facing role')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    add_bullet(doc, 'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub')
    add_bullet(doc, 'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution')

    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    add_bullet(doc, 'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 registered tool functions')
    add_bullet(doc, 'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1')

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 2. TOGETHERWORK — Client-Facing Technical Support Specialist
# ═══════════════════════════════════════════════════════════════
def build_togetherwork():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote experience supporting 150+ managed properties through ticket-based workflows, remote diagnostics, and direct client communication. 1,500+ hours of customer-facing troubleshooting resolving infrastructure and device issues. Experienced in documenting recurring issues, following structured escalation procedures, and building trusted client relationships through consistent follow-up. Google IT Support certified. Seeking a client-facing technical support role where empathetic communication and systematic troubleshooting improve the customer experience.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'Support & Troubleshooting', 'Device and infrastructure triage, remote diagnostics, ticket-based workflow management, structured escalation procedures, customer issue reproduction and documentation')
    add_skill_line(doc, 'Client Communication', 'Email, phone, and video-based support, proactive follow-up, non-technical explanations of technical concepts, client relationship management')
    add_skill_line(doc, 'Tools & Platforms', 'In-house ticketing platforms (Zendesk-style), Git, GitHub, Windows OS, Microsoft Office, Microsoft Teams')
    add_skill_line(doc, 'Technical', 'Python, SQL (coursework), HTML/CSS/JavaScript, API integration, CSV data handling')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    add_bullet(doc, 'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and device failures')
    add_bullet(doc, 'Deliver 1,500+ hours of 100% remote customer-facing support, guiding users through troubleshooting and resolution via email, phone, and screen-sharing')
    add_bullet(doc, 'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope')
    add_bullet(doc, 'Manage active tickets using an in-house platform, including outbound follow-up calls and status updates to maintain client trust')
    add_bullet(doc, 'Diagnose and triage issues with access points, switches, gateways, and modems, coordinating hardware replacement when needed')
    add_bullet(doc, 'Resolve an estimated 5\u201310 tickets per morning shift and 2\u20135 per night shift through remote troubleshooting and coordination')

    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Improved visibility and organization across dealership operations during a period of significant business growth')
    add_bullet(doc, 'Assisted walk-in clients, scheduled appointments, and answered phone lines in a client-facing administrative role')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    add_bullet(doc, 'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub')
    add_bullet(doc, 'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution')

    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    add_bullet(doc, 'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 tool functions')
    add_bullet(doc, 'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1')

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 3. USERPILOT — Implementation Specialist
# ═══════════════════════════════════════════════════════════════
def build_userpilot():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical professional who builds working systems and guides customers through implementation \u2014 from a spreadsheet-based inventory tracker that improved dealership operations, to a multi-agent AI assistant with 62 tool functions and a deterministic execution pipeline. 4+ years of remote support experience across 150+ properties with a focus on translating requirements into reliable solutions. Built and shipped a product with full lifecycle management including onboarding flow, UI design, and user documentation. Seeking an Implementation Specialist role where hands-on system configuration, customer engagement, and product-led thinking drive adoption.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'System Building', 'Pipeline architecture (PLAN/EXECUTE/VERIFY/REPORT), multi-agent design, API integration, process automation, onboarding flow design')
    add_skill_line(doc, 'Customer-Facing', 'Remote troubleshooting (1,500+ hours), customer guidance, cross-functional coordination, requirements gathering, training and enablement')
    add_skill_line(doc, 'Product & Analytics', 'User journey design, beta signup and engagement tracking, KPI-driven iteration, product landing page development')
    add_skill_line(doc, 'Languages & Frameworks', 'Python, React, TypeScript, HTML/CSS/JavaScript, SQL (coursework), Tailwind CSS')
    add_skill_line(doc, 'Tools', 'Git, GitHub, Vercel, Google Gemini API, OpenAI Whisper, NVIDIA CUDA, Windows/WSL')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    add_bullet(doc, 'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub (57 commits, 2 releases)')
    add_bullet(doc, 'Designed onboarding experience: first-run configuration wizard, setup flow, hotkey system, and floating UI for intuitive adoption')
    add_bullet(doc, 'Implemented smart model selection that auto-picks optimal speed/quality balance based on dictation length')
    add_bullet(doc, 'Built a product landing page using React, TypeScript, Tailwind CSS, and Vercel with working beta signup flow and interactive demo')
    add_bullet(doc, 'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution')

    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    add_bullet(doc, 'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 tool functions')
    add_bullet(doc, 'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1')
    add_bullet(doc, 'Built plugin chain builder with fallback resolution and config management for repeatable, reliable execution')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    add_bullet(doc, 'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and device failures')
    add_bullet(doc, 'Deliver 1,500+ hours of 100% remote customer-facing support, guiding users through troubleshooting and resolution')
    add_bullet(doc, 'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope')
    add_bullet(doc, 'Perform camera setup and implementation work alongside core network support responsibilities')

    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Improved visibility and organization across dealership operations during a period of significant business growth')

    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        'Relevant coursework: Database Design & Development, Business Management, Small Business Operations',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# 4. SENSI.AI — Technical Support Specialist
# ═══════════════════════════════════════════════════════════════
def build_sensi_ai():
    doc = create_doc()
    add_header_block(doc)
    add_spacer(doc)

    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote experience diagnosing device, network, and connectivity issues across 150+ managed properties. 1,500+ hours of direct B2B customer-facing support resolving infrastructure outages, equipment failures, and IoT device issues through structured diagnostic workflows. Experienced with network fundamentals including DHCP verification and device lifecycle management. Builds Python-based tools independently, including a shipped dictation app and a multi-agent AI system. Google IT Support certified. Seeking a Technical Support Specialist role at Sensi.AI where networking knowledge, device troubleshooting, and clear customer communication drive resolution.')

    add_spacer(doc)
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'Support & Networking', 'Device and infrastructure triage (access points, switches, gateways, modems), network connectivity diagnostics, DHCP verification, IoT device troubleshooting, remote diagnostics, equipment failure coordination')
    add_skill_line(doc, 'Operations', 'Ticket management, structured diagnostic workflows, escalation handling, outbound follow-up, issue documentation, cross-functional coordination')
    add_skill_line(doc, 'Tools & Platforms', 'In-house ticketing platforms, Git, GitHub, Windows OS, Microsoft Office, Microsoft Teams')
    add_skill_line(doc, 'Languages', 'Python, SQL (coursework), HTML/CSS/JavaScript, React, TypeScript')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')

    add_spacer(doc)
    add_section_header(doc, 'EXPERIENCE')
    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    add_bullet(doc, 'Provide remote technical support across 150+ managed properties, triaging issues with downed infrastructure, connectivity outages, and IoT device failures including access points, switches, gateways, and modems')
    add_bullet(doc, 'Deliver 1,500+ hours of 100% remote B2B customer-facing support, guiding property managers through troubleshooting and resolution via phone, email, and screen-sharing')
    add_bullet(doc, 'Perform network connectivity diagnostics including DHCP verification, device status checks, and service outage identification')
    add_bullet(doc, 'Follow structured diagnostic workflow: clarify scope, check for downed devices, walk through troubleshooting, reference past tickets, and escalate when outside scope')
    add_bullet(doc, 'Manage active tickets using an in-house platform, including outbound follow-up calls and status updates to minimize downtime')
    add_bullet(doc, 'Perform camera setup and implementation work including device configuration and connectivity verification')
    add_bullet(doc, 'Coordinate hardware replacement for failing equipment, managing device lifecycle from issue identification through resolution')

    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy')

    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Assisted walk-in clients, scheduled appointments, and answered phone lines in a client-facing role')

    add_spacer(doc)
    add_section_header(doc, 'PROJECTS')
    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    add_bullet(doc, 'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub')
    add_bullet(doc, 'Handled full product lifecycle: local ML inference, GPU optimization, UI design, installer packaging, license system, documentation, and beta distribution')

    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    add_bullet(doc, 'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 tool functions')
    add_bullet(doc, 'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1')

    add_spacer(doc)
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


# ═══════════════════════════════════════════════════════════════
# GENERATE ALL
# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    base = r'C:\Users\isaia\Projects\tools\job-project\04_Applications'

    builds = [
        ('Postman', 'Resume_Postman_TSE.docx', build_postman),
        ('Togetherwork_CRM', 'Resume_Togetherwork_TechSupport.docx', build_togetherwork),
        ('Userpilot', 'Resume_Userpilot_Implementation.docx', build_userpilot),
        ('Sensi_AI', 'Resume_Sensi_AI_TechSupport.docx', build_sensi_ai),
    ]

    for folder, filename, builder in builds:
        path = os.path.join(base, folder, filename)
        doc = builder()
        doc.save(path)
        print(f'Created: {path}')

    print('\nAll 4 Austin resumes generated.')
