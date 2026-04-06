"""Generate V2 Postman TSE resume .docx — keyword-optimized for Jobscan."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FONT = 'Arial'
NAME_SIZE = Pt(18)
CONTACT_SIZE = Pt(8)
LINK_SIZE = Pt(7.5)
SECTION_HEADER_SIZE = Pt(9.5)
BODY_SIZE = Pt(8.5)
TOP_MARGIN = 365760
BOTTOM_MARGIN = 365760
LEFT_MARGIN = 502920
RIGHT_MARGIN = 502920


def create_doc():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = TOP_MARGIN
    s.bottom_margin = BOTTOM_MARGIN
    s.left_margin = LEFT_MARGIN
    s.right_margin = RIGHT_MARGIN
    return doc

def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(1); pf.line_spacing = 1.0
    r = p.add_run(name); r.bold = True; r.font.size = NAME_SIZE; r.font.name = FONT

def add_contact_line(doc, text, is_link=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    r = p.add_run(text); r.font.size = LINK_SIZE if is_link else CONTACT_SIZE; r.font.name = FONT

def add_spacer(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.0

def add_section_header(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(4); pf.space_after = Pt(2); pf.line_spacing = 1.0
    r = p.add_run(text); r.bold = True; r.font.size = SECTION_HEADER_SIZE; r.font.name = FONT

def add_body_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(1); pf.line_spacing = 1.0
    r = p.add_run(text); r.font.size = BODY_SIZE; r.font.name = FONT

def add_skill_line(doc, category, content):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    r1 = p.add_run(f'{category}: '); r1.bold = True; r1.font.size = BODY_SIZE; r1.font.name = FONT
    r2 = p.add_run(content); r2.font.size = BODY_SIZE; r2.font.name = FONT

def add_job_header(doc, title_company, dates_location):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(3); pf.space_after = Pt(1); pf.line_spacing = 1.0
    r = p.add_run(title_company); r.bold = True; r.font.size = BODY_SIZE; r.font.name = FONT
    p2 = doc.add_paragraph()
    pf2 = p2.paragraph_format; pf2.space_before = Pt(0); pf2.space_after = Pt(1); pf2.line_spacing = 1.0
    r2 = p2.add_run(dates_location); r2.font.size = BODY_SIZE; r2.font.name = FONT

def add_project_header(doc, title, link=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(3); pf.space_after = Pt(0); pf.line_spacing = 1.0
    r = p.add_run(title); r.bold = True; r.font.size = BODY_SIZE; r.font.name = FONT
    if link:
        p2 = doc.add_paragraph()
        pf2 = p2.paragraph_format; pf2.space_before = Pt(0); pf2.space_after = Pt(0); pf2.line_spacing = 1.0
        r2 = p2.add_run(link); r2.font.size = LINK_SIZE; r2.font.name = FONT

def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    pf.left_indent = Inches(0.25)
    r = p.add_run(f'\u2022 {text}'); r.font.size = BODY_SIZE; r.font.name = FONT

def add_education(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        r = p.add_run(line); r.font.size = BODY_SIZE; r.font.name = FONT


def build():
    doc = create_doc()

    # Header
    add_name(doc, 'ISAIAH WASHINGTON')
    add_contact_line(doc, 'Cincinnati, OH | Open to Austin, TX and Remote')
    add_contact_line(doc, '513.544.9022 | isaiahwashington48@gmail.com')
    add_contact_line(doc, 'linkedin.com/in/isaiah-washington48 | github.com/Izayauh/whisper', is_link=True)
    add_spacer(doc)

    # Summary
    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    add_body_text(doc, 'Technical support professional with 4+ years of remote product support experience, researching and identifying solutions to resolve customer issues across 150+ managed properties. 1,500+ hours of customer-facing phone support, screen-sharing, and written communication \u2014 translating technical problems into clear guidance through analytical troubleshooting and structured diagnostic workflows. Builds API-driven internal tools independently using REST APIs and Python, including a shipped dictation app and a voice-controlled AI assistant with Google Gemini API function calling. Google IT Support certified. Eager to become a technical expert on the Postman platform, contributing to Postman\u2019s knowledge base and delivering consistent, high-quality product support to Postman\u2019s developer community.')
    add_spacer(doc)

    # Skills
    add_section_header(doc, 'TECHNICAL SKILLS')
    add_skill_line(doc, 'Product Support', 'Technical support and troubleshooting, phone support and screen-sharing, researching and identifying solutions, knowledge base contribution, customer issue documentation, structured escalation workflows')
    add_skill_line(doc, 'Infrastructure', 'Device and infrastructure triage (access points, switches, gateways, modems), remote diagnostics, networking fundamentals, connectivity issue resolution, equipment failure coordination')
    add_skill_line(doc, 'APIs & Integration', 'REST APIs, Google Gemini API (function calling, real-time audio streaming), OSC protocol, API development and testing, multi-agent tool dispatch')
    add_skill_line(doc, 'Tools & Platforms', 'Git, GitHub, Jira, in-house ticketing platforms, Postman (API testing), Windows OS, Microsoft Teams, Vercel')
    add_skill_line(doc, 'Languages', 'Python, JavaScript, SQL (coursework), HTML/CSS, React, TypeScript')
    add_skill_line(doc, 'Communication', 'Verbal communication skills, customer-facing technical guidance, cross-team collaboration, analytical problem-solving')
    add_skill_line(doc, 'Certification', 'Google IT Support \u2014 Coursera, 2022')
    add_spacer(doc)

    # Experience
    add_section_header(doc, 'EXPERIENCE')

    add_job_header(doc, 'Network Support Technician \u2014 Nexus WiFi', 'Oct 2021 \u2013 Present | West Chester, OH (100% Remote)')
    add_bullet(doc, 'Provide remote technical support and product support across 150+ managed properties, taking ownership of customer-reported issues and researching, diagnosing, and identifying solutions to resolve them')
    add_bullet(doc, 'Deliver 1,500+ hours of customer-facing phone support, screen-sharing, and written follow-ups, providing consistent verbal communication to both technical and non-technical users')
    add_bullet(doc, 'Lead analytical troubleshooting of access points, switches, gateways, and modems, coordinating with engineering teams to ensure speedy resolution of infrastructure failures')
    add_bullet(doc, 'Follow structured diagnostic workflow: clarify scope, research past tickets in the knowledge base, walk through troubleshooting steps, and escalate when outside scope')
    add_bullet(doc, 'Build and maintain internal tools and documentation to streamline support processes and improve resolution consistency')
    add_bullet(doc, 'Manage active tickets using an in-house platform, including outbound follow-up calls and status updates across support channels')
    add_bullet(doc, 'Resolve an estimated 5\u201310 tickets per morning shift and 2\u20135 per night shift, achieving consistent throughput through systematic troubleshooting')

    add_job_header(doc, 'Key Accounts Team Assistant \u2014 The Cincinnati Insurance Company', 'Feb 2023 \u2013 Aug 2023 | Fairfield, OH')
    add_bullet(doc, 'Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure consistent accuracy across cross-functional teams')

    add_job_header(doc, 'Administrative Assistant \u2014 Midway Auto Group', 'Jan 2019 \u2013 Dec 2023 | Middletown, OH')
    add_bullet(doc, 'Built and maintained an internal tracking tool (spreadsheet-based inventory system) covering cost, deliveries, and usage across the dealership\u2019s full vehicle stock')
    add_bullet(doc, 'Provided consistent phone support and walk-in client assistance, scheduling appointments and fielding product questions')
    add_spacer(doc)

    # Projects
    add_section_header(doc, 'PROJECTS')

    add_project_header(doc, 'Impulse \u2014 Privacy-First Dictation App', 'github.com/Izayauh/whisper | impulse-eight-lake.vercel.app')
    add_bullet(doc, 'Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python and OpenAI Whisper, published as a public beta on GitHub (57 commits, 2 releases)')
    add_bullet(doc, 'Handled full product lifecycle: researching model options, GPU optimization, UI design, installer packaging, license system, knowledge base documentation, and beta distribution')
    add_bullet(doc, 'Ran experiments with multiple Whisper model sizes to identify optimal speed/quality balance based on dictation length')

    add_project_header(doc, 'Live Pilot \u2014 Voice-Controlled AI DAW Assistant (Functional Prototype)')
    add_bullet(doc, 'Built a voice-controlled AI assistant for Ableton Live using Python, Google Gemini API, and OSC protocol with a 6-agent architecture and 62 registered tool functions called via REST APIs')
    add_bullet(doc, 'Designed a deterministic pipeline replacing iterative LLM loops, reducing API calls per operation from 30+ to 1 \u2014 demonstrating analytical problem-solving and technical expertise in API design')
    add_bullet(doc, 'Built internal tools for crash recovery, process lifecycle management, and idempotent operations for reliable execution')
    add_spacer(doc)

    # Education
    add_section_header(doc, 'EDUCATION')
    add_education(doc, [
        'Miami University \u2014 Cincinnati, OH',
        'Business Management Technology coursework',
        '67 credit hours toward Associate of Applied Business',
        'Attended through Spring 2024',
        '',
        'Google IT Support Certification \u2014 Coursera, November 2022',
    ])

    return doc


if __name__ == '__main__':
    doc = build()
    out = r'C:\Users\isaia\Projects\tools\job-project\04_Applications\Postman\Resume_Postman_TSE.docx'
    doc.save(out)
    print(f'Created: {out}')

    # Word count check
    count = 0
    for p in doc.paragraphs:
        count += len(p.text.split())
    print(f'Word count: {count}')
