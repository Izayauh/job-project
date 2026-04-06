import os
import glob
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT = 'Arial'
NAME_SIZE = Pt(18)
CONTACT_SIZE = Pt(8)
LINK_SIZE = Pt(7.5)
SECTION_HEADER_SIZE = Pt(9.5)
BODY_SIZE = Pt(8.5)
TOP_MARGIN = 365760
BOTTOM_MARGIN = 365760
LEFT_MARGIN = 502920
RIGHT_MARGIN = 502920

def create_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    return doc

def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(name)
    r.bold = True
    r.font.size = NAME_SIZE
    r.font.name = FONT

def add_contact_line(doc, text, is_link=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = LINK_SIZE if is_link else CONTACT_SIZE
    r.font.name = FONT

def add_spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = True
    r.font.size = SECTION_HEADER_SIZE
    r.font.name = FONT

def add_body_text(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = BODY_SIZE
    r.font.name = FONT

def add_skill_line(doc, line):
    if ':' in line:
        category, content = line.split(':', 1)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        r1 = p.add_run(f'{category}:')
        r1.bold = True
        r1.font.size = BODY_SIZE
        r1.font.name = FONT
        r2 = p.add_run(content)
        r2.font.size = BODY_SIZE
        r2.font.name = FONT
    else:
        add_body_text(doc, line, justify=False)

def add_job_header(doc, line):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    r = p.add_run(line)
    r.bold = True
    r.font.size = BODY_SIZE
    r.font.name = FONT

def add_job_subheader(doc, line):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    r = p.add_run(line)
    r.font.size = BODY_SIZE
    r.font.name = FONT

def add_bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.25)
    _text = text if text.startswith('\u2022') else f'\u2022 {text}'
    r = p.add_run(_text)
    r.font.size = BODY_SIZE
    r.font.name = FONT

def process_resume_md(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = [line.strip() for line in f.readlines()]
    
    # Extract only the text between ```text and ```
    text_content = []
    in_text_block = False
    for line in lines:
        if line.startswith('```text'):
            in_text_block = True
            continue
        elif line.startswith('```') and in_text_block:
            in_text_block = False
            continue
        if in_text_block:
            text_content.append(line)

    doc = create_doc()
    
    section = None
    header_idx = 0
    for line in text_content:
        if not line.strip():
            continue
            
        if header_idx == 0 and "ISAIAH WASHINGTON" in line:
            add_name(doc, line)
            header_idx += 1
            continue
        elif header_idx in (1, 2, 3):
            is_link = ("linkedin.com" in line.lower() or "github.com" in line.lower())
            add_contact_line(doc, line, is_link=is_link)
            header_idx += 1
            if header_idx == 4:
                add_spacer(doc)
            continue
            
        if line in ["SUMMARY", "SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION"]:
            section = line
            add_spacer(doc)
            add_section_header(doc, section)
            continue
            
        if section == "SUMMARY":
            add_body_text(doc, line)
        elif section == "SKILLS":
            add_skill_line(doc, line)
        elif section in ["EXPERIENCE", "PROJECTS"]:
            if line.startswith('\u2022') or line.startswith('- '):
                bullet_text = line.lstrip('\u2022- ')
                add_bullet(doc, bullet_text)
            elif '|' in line and not "github.com" in line:
                if 'Present' in line or ('2019' in line and '2023' in line) or ('2023' in line and 'Aug 2023' in line):
                    # Subheader
                    add_job_subheader(doc, line)
                else:
                    add_job_header(doc, line)
            elif "github.com" in line or "vercel.app" in line:
                add_job_subheader(doc, line)
            else:
                add_job_header(doc, line)
        elif section == "EDUCATION":
            add_body_text(doc, line, justify=False)

    doc.save(docx_path)

if __name__ == "__main__":
    companies = ["gWorks", "Postman", "Connecteam", "STRATATrust", "SensiAI", "Togetherwork", "Rippling"]
    base_dir = r"C:\Users\isaia\Projects\tools\job-project\04_Applications"
    
    for comp in companies:
        comp_dir = os.path.join(base_dir, comp)
        md_files = glob.glob(os.path.join(comp_dir, "Resume_*_v4.md"))
        if not md_files:
            continue
        md_path = md_files[0]
        
        # Create resume dir
        resume_dir = os.path.join(comp_dir, "resume")
        os.makedirs(resume_dir, exist_ok=True)
        
        docx_path = os.path.join(resume_dir, f"Resume.docx")
        
        try:
            process_resume_md(md_path, docx_path)
            print(f"Generated {docx_path}")
        except Exception as e:
            print(f"Failed to process {comp}: {e}")
