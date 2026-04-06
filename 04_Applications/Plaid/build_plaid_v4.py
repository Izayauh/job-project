"""Generate ATS-optimized PDF + DOCX resume for Plaid TSE v4.
Changes from v3:
- Negative words removed: 'failures' -> 'challenges', 'downed' -> 'disrupted',
  'urgent' -> 'high-priority', 'recurring issues' -> 'recurring patterns',
  'debugging' -> 'analyzing'
- Word count boosted toward 800+ with expanded bullets
- Outputs both PDF and DOCX
"""
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = os.path.dirname(__file__)

# ── Shared Resume Content ──────────────────────────────────────

SUMMARY = (
    "Technical Support Engineer with 4+ years of remote, customer-facing experience "
    "troubleshooting infrastructure and connectivity challenges across 150+ managed "
    "properties. Experienced in using ticketing systems, internal tools, and structured "
    "diagnostic workflows to investigate root cause, resolve technical issues, and "
    "escalate effectively. Collaborative team member who builds Python-based tooling "
    "independently, including coding projects that integrate REST APIs and multi-agent "
    "architectures. Customer-centric communicator with strong written and verbal "
    "communication skills who is empathetic to customer needs and genuinely passionate "
    "about delivering a world-class support experience. Eager to bring collaboration "
    "and technical depth to a financial services team helping customers resolve complex "
    "integration issues."
)

SKILLS = [
    ("Languages & Coding:", "Python, JavaScript, HTML, CSS, SQL (coursework)"),
    ("APIs & Integration:", "REST APIs, API debugging and testing, Google Gemini API (function calling), OSC protocol, structured request/response handling"),
    ("Ticketing & Tools:", "Ticketing systems (Jira, Zendesk), internal tools, browser developer tools, documentation platforms"),
    ("Tooling & Platforms:", "Git, GitHub, Vercel, Windows/WSL, Google Workspace"),
    ("Certification:", "Google IT Support - Coursera, 2022"),
]

NEXUS_BULLETS = [
    "Provide remote technical support across 150+ managed properties, troubleshooting disrupted infrastructure, connectivity challenges, and device performance across multiple environments while balancing priorities through effective time management",
    "Deliver 1,500+ hours of customer-facing support, taking an empathetic and customer-centric approach to understanding customer needs and guiding users through step-by-step troubleshooting and resolution",
    "Investigate root cause of connectivity challenges using internal tools and documentation, following structured diagnostic workflows to clarify scope, verify symptoms, and escalate through collaboration with appropriate teams when outside scope",
    "Manage 5-10 tickets per morning shift and 2-5 per night shift using ticketing systems, including detailed documentation, outbound follow-up calls, and status updates to ensure customer satisfaction",
    "Diagnose and triage technical support requests involving access points, switches, gateways, and modems, coordinating hardware replacement and working to collaborate cross-functionally with field technicians and vendors to reach resolution",
    "Work independently in a fully remote environment as a self-starter, proactively managing daily ticket queues and prioritizing high-priority customer requests while maintaining consistent collaboration with team members",
    "Contribute to process improvement by identifying recurring patterns, documenting solutions for future reference, refining troubleshooting workflows, and choosing to collaborate with teammates to share findings and improve outcomes",
]

CIC_BULLETS = [
    "Managed data input for key insured account policies in a financial services environment, coordinating with underwriters and assistants to ensure accuracy and meet customer needs",
    "Gained direct exposure to financial services operations, compliance workflows, and collaborative cross-functional coordination between teams",
]

MIDWAY_BULLETS = [
    "Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership's full vehicle stock, improving operational visibility",
    "Assisted walk-in clients, scheduled appointments, and answered phone lines in a customer-facing administrative role, developing strong written and verbal communication skills through daily collaboration",
]

IMPULSE_BULLETS = [
    "Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python coding and OpenAI Whisper, published as a public beta on GitHub (57 commits, 2 releases)",
    "Built a product landing page using JavaScript, HTML, CSS, React, TypeScript, and Tailwind CSS deployed on Vercel with working beta signup flow and interactive demo",
    "Handled full product lifecycle including tooling, documentation, UI design, installer packaging, and beta distribution through collaborative planning and execution",
    "Contributed to process improvement through brainstorming sessions on feature prioritization, user feedback integration, and iterative development",
]

JARVIS_BULLETS = [
    "Built a voice-controlled AI assistant for Ableton Live using Python coding, Google Gemini REST API, and OSC protocol with a 6-agent architecture and 62 tool functions",
    "Integrated REST API endpoints for real-time audio processing and function calling, with structured request/response handling, analyzing integration challenges, and collaborative iteration on design decisions",
    "Designed internal tooling to replace iterative LLM loops with a deterministic pipeline, reducing API calls per operation from 30+ to 1 through systematic testing and documentation",
]


# ── PDF Builder ────────────────────────────────────────────────

class ResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=20)

    def _c(self, t):
        return t.replace("\u2014", "-").replace("\u2013", "-").replace("\u2022", "-").replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')

    def section_header(self, text):
        self.set_font("Helvetica", "B", 11)
        self.cell(0, 7, self._c(text), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(60, 60, 60)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def body_text(self, text):
        self.set_font("Helvetica", "", 10)
        self.multi_cell(0, 5, self._c(text), new_x="LMARGIN", new_y="NEXT")

    def bullet(self, text):
        self.set_font("Helvetica", "", 10)
        self.cell(5, 5, "-", new_x="END")
        self.multi_cell(0, 5, self._c(text), new_x="LMARGIN", new_y="NEXT")

    def job_title(self, title, company):
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 6, self._c(f"{title} - {company}"), new_x="LMARGIN", new_y="NEXT")

    def job_meta(self, text):
        self.set_font("Helvetica", "I", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, self._c(text), new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)
        self.ln(2)

    def project_title(self, name, subtitle=""):
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 6, self._c(name), new_x="LMARGIN", new_y="NEXT")
        if subtitle:
            self.set_font("Helvetica", "I", 9)
            self.set_text_color(80, 80, 80)
            self.cell(0, 5, self._c(subtitle), new_x="LMARGIN", new_y="NEXT")
            self.set_text_color(0, 0, 0)
        self.ln(1)

    def skill_row(self, label, detail):
        self.set_font("Helvetica", "B", 9.5)
        self.cell(40, 5, self._c(label), new_x="END")
        self.set_font("Helvetica", "", 9.5)
        self.multi_cell(0, 5, self._c(detail), new_x="LMARGIN", new_y="NEXT")
        self.ln(1)


def build_pdf():
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 15, 18)

    # Header
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, "ISAIAH WASHINGTON", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 5, pdf._c("Cincinnati, OH (Eastern Time)  |  Remote  |  513.544.9022  |  isaiahwashington48@gmail.com"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, pdf._c("linkedin.com/in/isaiah-washington48  |  github.com/Izayauh/whisper"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    pdf.section_header("PROFESSIONAL SUMMARY")
    pdf.body_text(SUMMARY)
    pdf.ln(3)

    pdf.section_header("TECHNICAL SKILLS")
    for label, detail in SKILLS:
        pdf.skill_row(label, detail)
    pdf.ln(2)

    pdf.section_header("EXPERIENCE")
    pdf.job_title("Network Support Technician", "Nexus WiFi")
    pdf.job_meta("Oct 2021 - Present  |  West Chester, OH (100% Remote)")
    for b in NEXUS_BULLETS:
        pdf.bullet(b)
    pdf.ln(2)

    pdf.job_title("Key Accounts Team Assistant", "The Cincinnati Insurance Company")
    pdf.job_meta("Feb 2023 - Aug 2023  |  Fairfield, OH (Financial Services)")
    for b in CIC_BULLETS:
        pdf.bullet(b)
    pdf.ln(2)

    pdf.job_title("Administrative Assistant", "Midway Auto Group")
    pdf.job_meta("Jan 2019 - Dec 2023  |  Middletown, OH")
    for b in MIDWAY_BULLETS:
        pdf.bullet(b)
    pdf.ln(3)

    pdf.section_header("TECHNICAL PROJECTS")
    pdf.project_title("Impulse - Privacy-First Dictation App", "github.com/Izayauh/whisper  |  impulse-eight-lake.vercel.app")
    for b in IMPULSE_BULLETS:
        pdf.bullet(b)
    pdf.ln(2)

    pdf.project_title("Live Pilot - Voice-Controlled AI DAW Assistant", "(Functional Prototype)")
    for b in JARVIS_BULLETS:
        pdf.bullet(b)
    pdf.ln(3)

    pdf.section_header("EDUCATION")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "Miami University - Cincinnati, OH", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, "Business Management Technology coursework - 67 credit hours toward Associate of Applied Business", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 9.5)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, "Attended through Spring 2024", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9.5)
    pdf.cell(0, 5, "Relevant coursework: Database Design & Development (SQL), HTML/CSS/JavaScript", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "Google IT Support Certification - Coursera, November 2022", new_x="LMARGIN", new_y="NEXT")

    out = os.path.join(BASE, "Resume_Plaid_TSE_v4.pdf")
    pdf.output(out)
    print(f"PDF: {out}")
    return out


# ── DOCX Builder ───────────────────────────────────────────────

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Add bottom border via paragraph formatting
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '3C3C3C',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

def add_job(doc, title, company, meta):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(0)
    run = p.add_run(f"{title} - {company}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(2)
    run2 = p2.add_run(meta)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(100, 100, 100)

def add_project(doc, name, subtitle=""):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(0)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    if subtitle:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.name = "Calibri"
        run2.font.color.rgb = RGBColor(80, 80, 80)

def add_skill_line(doc, label, detail):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run_label = p.add_run(label + " ")
    run_label.bold = True
    run_label.font.size = Pt(9.5)
    run_label.font.name = "Calibri"
    run_detail = p.add_run(detail)
    run_detail.font.size = Pt(9.5)
    run_detail.font.name = "Calibri"


def build_docx():
    doc = Document()

    # Adjust margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header - Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.space_after = Pt(0)
    run = name_p.add_run("ISAIAH WASHINGTON")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"

    # Contact line 1
    contact1 = doc.add_paragraph()
    contact1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact1.space_before = Pt(0)
    contact1.space_after = Pt(0)
    r = contact1.add_run("Cincinnati, OH (Eastern Time)  |  Remote  |  513.544.9022  |  isaiahwashington48@gmail.com")
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(80, 80, 80)

    # Contact line 2
    contact2 = doc.add_paragraph()
    contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact2.space_before = Pt(0)
    contact2.space_after = Pt(6)
    r = contact2.add_run("linkedin.com/in/isaiah-washington48  |  github.com/Izayauh/whisper")
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(80, 80, 80)

    # Summary
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(SUMMARY)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    # Skills
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, detail in SKILLS:
        add_skill_line(doc, label, detail)

    # Experience
    add_section_heading(doc, "EXPERIENCE")

    add_job(doc, "Network Support Technician", "Nexus WiFi", "Oct 2021 - Present  |  West Chester, OH (100% Remote)")
    for b in NEXUS_BULLETS:
        add_bullet(doc, b)

    add_job(doc, "Key Accounts Team Assistant", "The Cincinnati Insurance Company", "Feb 2023 - Aug 2023  |  Fairfield, OH (Financial Services)")
    for b in CIC_BULLETS:
        add_bullet(doc, b)

    add_job(doc, "Administrative Assistant", "Midway Auto Group", "Jan 2019 - Dec 2023  |  Middletown, OH")
    for b in MIDWAY_BULLETS:
        add_bullet(doc, b)

    # Projects
    add_section_heading(doc, "TECHNICAL PROJECTS")

    add_project(doc, "Impulse - Privacy-First Dictation App", "github.com/Izayauh/whisper  |  impulse-eight-lake.vercel.app")
    for b in IMPULSE_BULLETS:
        add_bullet(doc, b)

    add_project(doc, "Live Pilot - Voice-Controlled AI DAW Assistant", "(Functional Prototype)")
    for b in JARVIS_BULLETS:
        add_bullet(doc, b)

    # Education
    add_section_heading(doc, "EDUCATION")

    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(0)
    run = p.add_run("Miami University - Cincinnati, OH")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(0)
    run = p2.add_run("Business Management Technology coursework - 67 credit hours toward Associate of Applied Business")
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.space_before = Pt(0)
    p3.space_after = Pt(0)
    run = p3.add_run("Attended through Spring 2024")
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(100, 100, 100)

    p4 = doc.add_paragraph()
    p4.space_before = Pt(0)
    p4.space_after = Pt(6)
    run = p4.add_run("Relevant coursework: Database Design & Development (SQL), HTML/CSS/JavaScript")
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"

    p5 = doc.add_paragraph()
    p5.space_before = Pt(4)
    run = p5.add_run("Google IT Support Certification - Coursera, November 2022")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    out = os.path.join(BASE, "Resume_Plaid_TSE_v4.docx")
    doc.save(out)
    print(f"DOCX: {out}")
    return out


if __name__ == "__main__":
    # Count words
    all_text = SUMMARY
    for _, d in SKILLS:
        all_text += " " + d
    for bullets in [NEXUS_BULLETS, CIC_BULLETS, MIDWAY_BULLETS, IMPULSE_BULLETS, JARVIS_BULLETS]:
        for b in bullets:
            all_text += " " + b
    wc = len(all_text.split())
    print(f"Approx word count (content only): {wc}")

    build_pdf()
    build_docx()
    print("\nDone - both formats generated.")
