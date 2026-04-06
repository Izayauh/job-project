"""Generate ATS-optimized PDF + DOCX resume for Turvo CSE."""
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = os.path.dirname(__file__)

# ── Shared Resume Content ──────────────────────────────────────

SUMMARY = (
    "Customer Support Engineer with 4+ years of customer-centric application support "
    "experience performing root cause analysis, troubleshooting, and escalation handling "
    "across 150+ managed properties in a fast-paced environment. 1,500+ hours of "
    "customer-facing technical support ensuring resolution within service level agreements "
    "(SLAs) while collaborating with site reliability engineering and product teams. Builds "
    "Python scripting tools independently, including REST API / OpenAPI integration work and "
    "structured request/response debugging. Passionate about supply chain and Transportation "
    "Management System (TMS) technology, delivering outstanding customer experience through "
    "collaboration, adaptability, and excellent written and verbal communication skills."
)

SKILLS = [
    ("Troubleshooting:", "Root cause analysis, structured diagnostic workflows, device and infrastructure triage, third-party integration diagnostics, connectivity resolution, escalation coordination, Java debugging fundamentals"),
    ("Programming & APIs:", "Python scripting, REST APIs / OpenAPI, Google Gemini API (function calling), MySQL / SQL (coursework), browser developer tools, structured request/response handling"),
    ("Documentation:", "Technical documentation, knowledge base articles, ticket documentation, structured follow-up write-ups, customer-facing correspondence"),
    ("Support Tools:", "CRM systems, remote support tools, analytics dashboards, ticketing systems, internal diagnostic tools, Google Workspace"),
    ("Platforms:", "Git, GitHub, Vercel, Windows/WSL"),
    ("Certification:", "Google IT Support - Coursera, 2022"),
]

NEXUS_BULLETS = [
    "Provide remote application support and technical support across 150+ managed properties serving logistics and freight operations, diagnosing and resolving connectivity challenges and device performance issues in a fast-paced environment",
    "Deliver 1,500+ hours of customer-centric support with consistent reliability, ensuring resolution within service level agreements (SLAs) while guiding users through troubleshooting and resolution to drive customer success",
    "Perform root cause analysis on connectivity challenges using structured diagnostics and technical documentation, escalating to collaborate with site reliability and engineering teams when outside scope",
    "Work independently to manage 5-10 tickets per morning shift and 2-5 per night shift, ensuring resolution with detailed documentation, outbound follow-up calls, and status updates",
    "Diagnose and resolve product and third-party integration issues with access points, switches, gateways, and modems, coordinating hardware replacement through collaborative cross-functional coordination",
    "Create and maintain technical documentation including knowledge base articles, ticket write-ups, and customer-facing correspondence for internal teams and customers",
    "Identify recurring issues and performance improvement opportunities, sharing customer insights to influence product enhancements and drive continuous improvement",
    "Act as the voice of the customer by relaying feedback and patterns to management, advocating for improved workflows and customer experience across the support team",
]

CIC_BULLETS = [
    "Managed data input for key insured account policies, coordinating with underwriters and assistants to ensure accuracy and meet customer needs in a detail-oriented, fast-paced environment",
    "Gained direct exposure to cross-functional coordination between teams, strengthening collaboration and written and verbal communication skills",
]

MIDWAY_BULLETS = [
    "Built and maintained a spreadsheet-based inventory tracking system covering cost, deliveries, and usage across the dealership's full vehicle stock, improving operational visibility and analytics through adaptability to changing business needs",
    "Assisted walk-in clients, scheduled appointments, and answered phone lines in a customer-facing administrative role, strengthening customer experience instincts through daily collaboration",
]

IMPULSE_BULLETS = [
    "Built and shipped a GPU-accelerated speech-to-text dictation app for Windows using Python scripting and OpenAI Whisper, published as a public beta on GitHub (57 commits, 2 releases)",
    "Handled full product lifecycle including technical documentation, knowledge base creation, UI design, installer packaging, and beta distribution through collaborative planning and execution",
    "Contributed to continuous improvement through brainstorming sessions on feature prioritization, user feedback integration, and iterative development",
]

JARVIS_BULLETS = [
    "Built a voice-controlled AI assistant for Ableton Live using Python scripting, Google Gemini REST API, and OSC protocol with a 6-agent architecture and 62 tool functions",
    "Integrated REST API / OpenAPI endpoints for real-time audio processing and function calling, with structured request/response handling and analysis of third-party integration issues",
    "Designed internal tooling to replace iterative LLM loops with a deterministic pipeline, reducing API calls per operation from 30+ to 1 through systematic testing and technical documentation",
]


# ── PDF Builder ────────────────────────────────────────────────

class ResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=20)

    def _c(self, t):
        return t.replace("\u2014", "-").replace("\u2013", "-").replace("\u2022", "-").replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')

    def section_header(self, text):
        self.set_font("Helvetica", "B", 11)
        self.cell(0, 7, self._c(text), new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(60, 60, 60)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def body_text(self, text):
        self.set_font("Helvetica", "", 10)
        self.multi_cell(0, 5, self._c(text), new_x="LMARGIN", new_y="NEXT")

    def bullet(self, text):
        self.set_font("Helvetica", "", 10)
        self.cell(5, 5, "-", new_x="END")
        self.multi_cell(0, 5, self._c(text), new_x="LMARGIN", new_y="NEXT")

    def job_title(self, title, company):
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 6, self._c(f"{title} - {company}"), new_x="LMARGIN", new_y="NEXT")

    def job_meta(self, text):
        self.set_font("Helvetica", "I", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, self._c(text), new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)
        self.ln(2)

    def project_title(self, name, subtitle=""):
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 6, self._c(name), new_x="LMARGIN", new_y="NEXT")
        if subtitle:
            self.set_font("Helvetica", "I", 9)
            self.set_text_color(80, 80, 80)
            self.cell(0, 5, self._c(subtitle), new_x="LMARGIN", new_y="NEXT")
            self.set_text_color(0, 0, 0)
        self.ln(1)

    def skill_row(self, label, detail):
        self.set_font("Helvetica", "B", 9.5)
        self.cell(40, 5, self._c(label), new_x="END")
        self.set_font("Helvetica", "", 9.5)
        self.multi_cell(0, 5, self._c(detail), new_x="LMARGIN", new_y="NEXT")
        self.ln(1)


def build_pdf():
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 15, 18)

    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, "ISAIAH WASHINGTON", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 5, pdf._c("Cincinnati, OH  |  Open to Dallas, TX and Remote  |  513.544.9022  |  isaiahwashington48@gmail.com"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, pdf._c("linkedin.com/in/isaiah-washington48  |  github.com/Izayauh/whisper"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    pdf.section_header("PROFESSIONAL SUMMARY")
    pdf.body_text(SUMMARY)
    pdf.ln(3)

    pdf.section_header("TECHNICAL SKILLS")
    for label, detail in SKILLS:
        pdf.skill_row(label, detail)
    pdf.ln(2)

    pdf.section_header("EXPERIENCE")
    pdf.job_title("Network Support Technician", "Nexus WiFi")
    pdf.job_meta("Oct 2021 - Present  |  West Chester, OH (100% Remote)")
    for b in NEXUS_BULLETS:
        pdf.bullet(b)
    pdf.ln(2)

    pdf.job_title("Key Accounts Team Assistant", "The Cincinnati Insurance Company")
    pdf.job_meta("Feb 2023 - Aug 2023  |  Fairfield, OH")
    for b in CIC_BULLETS:
        pdf.bullet(b)
    pdf.ln(2)

    pdf.job_title("Administrative Assistant", "Midway Auto Group")
    pdf.job_meta("Jan 2019 - Dec 2023  |  Middletown, OH")
    for b in MIDWAY_BULLETS:
        pdf.bullet(b)
    pdf.ln(3)

    pdf.section_header("TECHNICAL PROJECTS")
    pdf.project_title("Impulse - Privacy-First Dictation App", "github.com/Izayauh/whisper  |  impulse-eight-lake.vercel.app")
    for b in IMPULSE_BULLETS:
        pdf.bullet(b)
    pdf.ln(2)

    pdf.project_title("Live Pilot - Voice-Controlled AI DAW Assistant", "(Functional Prototype)")
    for b in JARVIS_BULLETS:
        pdf.bullet(b)
    pdf.ln(3)

    pdf.section_header("EDUCATION")
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "Miami University - Cincinnati, OH", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, "Business Management Technology coursework - 67 credit hours toward Associate of Applied Business", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "I", 9.5)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, "Attended through Spring 2024", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9.5)
    pdf.cell(0, 5, "Relevant coursework: Database Design & Development (MySQL / SQL),", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, "Information Technology fundamentals, Computer Science principles, Workplace Writing", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "Google IT Support Certification - Coursera, November 2022", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9.5)
    pdf.cell(0, 5, "Covers computer information systems, networking, troubleshooting, and IT support fundamentals", new_x="LMARGIN", new_y="NEXT")

    out = os.path.join(BASE, "Resume_Turvo_CSE.pdf")
    pdf.output(out)
    print(f"PDF: {out}")


# ── DOCX Builder ───────────────────────────────────────────────

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '3C3C3C',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

def add_job(doc, title, company, meta):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(0)
    run = p.add_run(f"{title} - {company}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(2)
    run2 = p2.add_run(meta)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(100, 100, 100)

def add_project(doc, name, subtitle=""):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(0)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    if subtitle:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.name = "Calibri"
        run2.font.color.rgb = RGBColor(80, 80, 80)

def add_skill_line(doc, label, detail):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.name = "Calibri"
    r2 = p.add_run(detail)
    r2.font.size = Pt(9.5)
    r2.font.name = "Calibri"


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.space_after = Pt(0)
    run = name_p.add_run("ISAIAH WASHINGTON")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"

    c1 = doc.add_paragraph()
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.space_before = Pt(0)
    c1.space_after = Pt(0)
    r = c1.add_run("Cincinnati, OH  |  Open to Dallas, TX and Remote  |  513.544.9022  |  isaiahwashington48@gmail.com")
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(80, 80, 80)

    c2 = doc.add_paragraph()
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.space_before = Pt(0)
    c2.space_after = Pt(6)
    r = c2.add_run("linkedin.com/in/isaiah-washington48  |  github.com/Izayauh/whisper")
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(80, 80, 80)

    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(SUMMARY)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, detail in SKILLS:
        add_skill_line(doc, label, detail)

    add_section_heading(doc, "EXPERIENCE")
    add_job(doc, "Network Support Technician", "Nexus WiFi", "Oct 2021 - Present  |  West Chester, OH (100% Remote)")
    for b in NEXUS_BULLETS:
        add_bullet(doc, b)
    add_job(doc, "Key Accounts Team Assistant", "The Cincinnati Insurance Company", "Feb 2023 - Aug 2023  |  Fairfield, OH")
    for b in CIC_BULLETS:
        add_bullet(doc, b)
    add_job(doc, "Administrative Assistant", "Midway Auto Group", "Jan 2019 - Dec 2023  |  Middletown, OH")
    for b in MIDWAY_BULLETS:
        add_bullet(doc, b)

    add_section_heading(doc, "TECHNICAL PROJECTS")
    add_project(doc, "Impulse - Privacy-First Dictation App", "github.com/Izayauh/whisper  |  impulse-eight-lake.vercel.app")
    for b in IMPULSE_BULLETS:
        add_bullet(doc, b)
    add_project(doc, "Live Pilot - Voice-Controlled AI DAW Assistant", "(Functional Prototype)")
    for b in JARVIS_BULLETS:
        add_bullet(doc, b)

    add_section_heading(doc, "EDUCATION")
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(0)
    run = p.add_run("Miami University - Cincinnati, OH")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(0)
    run = p2.add_run("Business Management Technology coursework - 67 credit hours toward Associate of Applied Business")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p3 = doc.add_paragraph()
    p3.space_before = Pt(0)
    p3.space_after = Pt(0)
    run = p3.add_run("Attended through Spring 2024")
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(100, 100, 100)
    p4 = doc.add_paragraph()
    p4.space_before = Pt(0)
    p4.space_after = Pt(0)
    run = p4.add_run("Relevant coursework: Database Design & Development (MySQL / SQL), Information Technology fundamentals, Computer Science principles, Workplace Writing")
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    p5 = doc.add_paragraph()
    p5.space_before = Pt(4)
    p5.space_after = Pt(0)
    run = p5.add_run("Google IT Support Certification - Coursera, November 2022")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p6 = doc.add_paragraph()
    p6.space_before = Pt(0)
    run = p6.add_run("Covers computer information systems, networking, troubleshooting, and IT support fundamentals")
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"

    out = os.path.join(BASE, "Resume_Turvo_CSE.docx")
    doc.save(out)
    print(f"DOCX: {out}")


if __name__ == "__main__":
    all_text = SUMMARY
    for _, d in SKILLS:
        all_text += " " + d
    for bullets in [NEXUS_BULLETS, CIC_BULLETS, MIDWAY_BULLETS, IMPULSE_BULLETS, JARVIS_BULLETS]:
        for b in bullets:
            all_text += " " + b
    print(f"Approx word count (content only): {len(all_text.split())}")
    build_pdf()
    build_docx()
    print("\nDone - both formats generated.")
